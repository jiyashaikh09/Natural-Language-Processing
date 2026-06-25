# Import the requires Libraries
import PyPDF2
from docx import Document
from progressbar import progressbar
from spire.pdf import PdfDocument, FileFormat as PdfFIleFormat
from spire.doc import FileFormat as docFormat
from spire.doc import Document as docDocument
from spire.doc import XHTMLValidationType
import argostranslate.package, argostranslate.translate
import translatehtml
from bs4 import BeautifulSoup
from pyhtml2pdf import converter
import threading
from tkinter import *
from tkinter.ttk import Style
import tkinter as tk
import tkinter.messagebox as msgbox
import openpyxl
import pylab as p
from pypdf import PdfReader
import os
from PyPDF2 import PdfMerger
from PIL import Image, ImageTk
from tkinter import filedialog, messagebox, scrolledtext, font, ttk
from tkinter.filedialog import askopenfile
import argostranslate.package
import argostranslate.translate
from fpdf import FPDF
from tkinter.scrolledtext import ScrolledText
#import LoginPage
import customtkinter
from customtkinter import CTk
import time
import datetime
import pathlib
from lxml import etree
#from Argospdftrans import text
import os


try:
    # python 2.x
    import Tkinter as tk
except ImportError:
    # python 3.x
    import tkinter as tk
from tkinter import messagebox
from openpyxl import Workbook
import openpyxl
import docx
from PIL import Image
# Function to set up proxy if available
import os
import urllib.request
import urllib.error



def test_connection(test_url, timeout=1):
    try:
        response = urllib.request.urlopen(test_url, timeout=timeout)
        if response.status == 200:
            return True
    except urllib.error.URLError as e:
        print(f"Connection test failed: {e.reason}")
    return False


def detect_proxy():
    test_url = "http://httpbin.org/ip"
    # try:
    # response = urllib.request.urlopen(test_url, timeout=2)
    if test_connection(test_url):
        print("Connected successfully. No proxy detected or proxy is functioning correctly.")
        return False  # No proxy detected or proxy is functioning correctly
    # except urllib.error.URLError as e:
    print(f"Initial connection test failed. Proxy might be in use.")
    return True  # Proxy might be blocking or not functioning correctly


def setup_proxy():
    proxy_address = 'proxy01.vodafone-is.de'
    proxy_port = '8080'
    proxy_url = f'http://{proxy_address}:{proxy_port}'

    # Set environment variables for proxy
    os.environ['http_proxy'] = proxy_url
    os.environ['https_proxy'] = proxy_url

    # Create proxy handler
    proxy_support = urllib.request.ProxyHandler({
        'http': proxy_url,
        'https': proxy_url
    })
    opener = urllib.request.build_opener(proxy_support)

    try:
        response = opener.open("http://httpbin.org/ip", timeout=1)
        if response.status == 200:
            print("Proxy setup successful")
            urllib.request.install_opener(opener)
        else:
            print("Proxy setup failed")
            raise urllib.error.URLError("Failed to reach test URL")
    except urllib.error.URLError as e:
        print(f"Proxy setup failed: {e.reason}")
        # Fallback to default opener
        opener = urllib.request.build_opener()
        urllib.request.install_opener(opener)


def remove_proxy():
    # Remove proxy environment variables
    os.environ.pop('http_proxy', None)
    os.environ.pop('https_proxy', None)

    # Create default opener
    opener = urllib.request.build_opener()
    urllib.request.install_opener(opener)

def get_current_directory_path():
    cwd = os.getcwd()
    return cwd

def get_current_folder_path():
    print("File      Path:", pathlib.Path().resolve())
    cwf = pathlib.Path().resolve()
    return cwf

def configure_network():
    try:
        if detect_proxy():
            print("Proxy detected. Configuring proxy settings...")
            setup_proxy()
        else:
            print("No proxy detected. Proceeding without proxy.")
            remove_proxy()
    except Exception as e:
        print(f"Error in network configuration: {e}")
        remove_proxy()


configure_network()


def example_function():
    try:
        response = urllib.request.urlopen("http://httpbin.org/ip")
        print(response.read().decode())
    except Exception as e:
        print(f"Error during network request: {e}")


from PIL import Image

# root = CTk()
root = tk.Tk()
# translate = 0
# Set the geometry of Tkinter frame
root.geometry("800x800")
# root.geometry("2700x1200")


# Create a canvas
canvas = Canvas(root, width=400, height=250, bg='#55a2aa', bd=0, relief='ridge', highlightthickness=0)
canvas.grid(columnspan=5, rowspan=5)

# frame

frame = Frame(root, borderwidth=0)
Label(frame, text="                                       ", fg='#55a2aa', bg='#55a2aa',
      font="Times 33 bold italic").pack(side=TOP)
frame.grid(row=0, column=8, columnspan=50)

root.iconbitmap('vois.png')
root.title('Language Translator(BQA/Testing)')
# screenwidth = root.winfo_screenwidth()
# screenwidth = root.winfo_screenwidth()
# screenheight = root.winfo_screenheight()
root['bg'] = '#55a2aa'
# Load Image
photo = Image.open('NewUI.png')
# Resized Image
resized_photo = photo.resize((370, 400))
photo = ImageTk.PhotoImage(resized_photo)
w = tk.Canvas(root, width=370, height=400, bg='#55a2aa', borderwidth=0)
w.grid(column=0, row=4, columnspan=9)
w.create_image(1, 0, anchor=tk.NW, image=photo)
w.grid_propagate(False)

my_text = tk.Canvas(root, height=25, width=70)
my_text = scrolledtext.ScrolledText(root, wrap=tk.NONE, height=25, width=70)
my_text.grid(column=10, row=4, columnspan=10)

# Create horizontal scrollbar
h_scrollbar = tk.Scrollbar(root, orient=tk.HORIZONTAL, command=my_text.xview)
h_scrollbar.grid(row=3, column=10, columnspan=12, sticky='W' + 'E' + 'S')
# h_scrollbar.grid(side=tk.BOTTOM, fill=tk.X)
my_text['xscrollcommand'] = h_scrollbar.set
# Instruction
label = Label(root, text="Select the file from your system", height=2, width=40, fg="black",
              bg='#77B5BB')
label.grid(row=12, column=1, columnspan=5)
# root.geometry("300x200")
root.geometry("1000x800")


# Translate German to English Function

def argo_function(textbox_data, from_code, to_code, progressbar):
    argostranslate.package.update_package_index()
    available_packages = argostranslate.package.get_available_packages()
    package_to_install = next(
        filter(
            lambda x: x.from_code == from_code and x.to_code == to_code, available_packages
        )
    )
    argostranslate.package.install_from_path(package_to_install.download())

    # Translate
    translatedText = argostranslate.translate.translate(textbox_data, from_code, to_code)
    # print(translatedText)
    my_text.delete(1.0, END)
    # browse_text.set("Select File")
    my_text.insert(1.0, translatedText)
    # print("CLosing the Progress Bar")
    #print("In Thread translate stop")
    progressbar.stop()
    #print(("CLosing the progressbar"))
    progressbar.destroy()

    # return translatedText

def translate_german_to_english(input_file,typeofFile):
    try:
        print("In Translation function",input_file)
        dict_valuesReplace = {
            "Single price": "Unit price",
            "KUNDEN NR": "Customer No.",
            "USt.": "Vat",
            "RECOVERY": "YOUR BILL",
            "Please worry for sufficient account coverage.": "Please ensure that your account has sufficient funds.",
            "Woman": "Mrs.",
            "www.vodafone.de/Calculation":"www.vodafone.de/rechnung"
        }
        # Load the input DOCX file
        doc = Document(input_file)

        # Check if the translation package is already installed
        installed_languages = argostranslate.translate.get_installed_languages()
        from_lang = next((lang for lang in installed_languages if lang.code == "de"), None)
        to_lang = next((lang for lang in installed_languages if lang.code == "en"), None)

        if not from_lang or not to_lang:
            # Download and install Argos Translate package if not installed
            available_packages = argostranslate.package.get_available_packages()
            available_package = next(
                (pkg for pkg in available_packages if pkg.from_code == 'de' and pkg.to_code == 'en'), None
            )
            if available_package:
                download_path = available_package.download()
                argostranslate.package.install_from_path(download_path)
                installed_languages = argostranslate.translate.get_installed_languages()
                from_lang = next(lang for lang in installed_languages if lang.code == "de")
                to_lang = next(lang for lang in installed_languages if lang.code == "en")

        def translate_text(text, from_lang, to_lang):
            ''' Function to translate word file using source and translation language'''
            return from_lang.get_translation(to_lang).translate(text)
            # Translate paragraphs while preserving formatting

        for para in doc.paragraphs:
            for run in para.runs:

                for key, value in dict_valuesReplace.items():
                    run.text = run.text.replace(key, value)
                if "Evaluation Warning" not in run.text:
                    run.text = translate_text(run.text, from_lang, to_lang)
                else:
                    run.text = run.text.replace("Evaluation Warning*", "")
                    print("Got Evaluation Warning in para", run.text)

                if "www" not in run.text:
                    run.text = translate_text(run.text,from_lang, to_lang)
                else:

                    if run.text == "www.vodafone.de/contact":
                        run.text = run.text.replace("www.vodafone.de/contact","www.vodafone.de/kontakt")
                    if run.text == "www.vodafone.de/frequency conversion":
                        run.text = run.text.replace("www.vodafone.de/frequency conversion", "www.vodafone.de/frequenzumbelegung")
                    if run.text == "www.vodafone.de/Calculation":
                        run.text = run.text.replace("www.vodafone.de/Calculation","www.vodafone.de/rechnung")
                        #print("Found", run.text)
                    #print("Got www in para1",run.text)

            # Translate tables while preserving formatting
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            for key, value in dict_valuesReplace.items():
                                run.text = run.text.replace(key, value)

                            if "www" not in run.text:
                                run.text = translate_text(run.text,from_lang,to_lang)
                            else:
                                print("Got www in table",run.text)
                            if "Evaluation Warning" not in run.text:
                                run.text = translate_text(run.text, from_lang, to_lang)
                            else:
                                run.text = run.text.replace("Evaluation Warning*", "")
                                print("Got Evaluation Warning in table", run.text)
                            if run.text == "www.vodafone.de/contact":
                                run.text = run.text.replace("www.vodafone.de/contact", "www.vodafone.de/kontakt")
                            if "133rd" in run.text:
                                run.text = run.text.replace("133rd", "")
                                #print("133rd", run.text)

            # Translate headers and footers while preserving formatting
        for section in doc.sections:
            header = section.header
            for para in header.paragraphs:
                for run in para.runs:
                    for key, value in dict_valuesReplace.items():
                        run.text = run.text.replace(key, value)
                    if run.text == "www.vodafone.de/Calculation":
                        print("Found", run.text)
                    if "www." not in run.text:
                        run.text = translate_text(run.text,from_lang, to_lang)
                        #print("Para in headers",type(run.text))
                    else:
                        print("Got www in header",run.text)

                    if "Evaluation Warning" not in run.text:
                        run.text = translate_text(run.text, from_lang, to_lang)
                    else:
                        run.text = run.text.replace("Evaluation Warning*", "")
                        print("Got Evaluation Warning in header", run.text)
                    if "133rd" in run.text:
                        run.text = run.text.replace("133rd","")
                        #print("133rd", run.text)
            footer = section.footer
            for para in footer.paragraphs:
                for run in para.runs:
                    for key, value in dict_valuesReplace.items():
                        run.text = run.text.replace(key, value)
                    if run.text == "www.vodafone.de/Calculation":
                        print("Found", run.text)
                    if "www." not in run.text:
                        run.text = translate_text(run.text,from_lang, to_lang)
                        #print("Para in headers", type(run.text))
                    else:
                        print("Got www in footer",run.text)
                    if "Evaluation Warning" not in run.text:
                        run.text = translate_text(run.text, from_lang, to_lang)
                    else:
                        run.text = run.text("Evaluation Warning*", "")
                        print("Got Evaluation Warning in footer", run.text)
        # Save the translated document to the output file
        op_path = str(get_current_folder_path())
        #print(op_path)
        if typeofFile=='word':
            output_file = op_path +"\\Result_"+getcurrentDateTime()+".docx"
        else:
            output_file = op_path + "\\output\\Result_" + getcurrentDateTime() + ".docx"

        doc.save(output_file)
        #messagebox.showinfo("Successful...")
        messagebox.showinfo('Successfully',"Document has saved Successfully...")

        # Print a success message
        print(f"The German document '{input_file}' was successfully translated to English and saved as '{output_file}'")

    except Exception as e:
        print(f"An error occurred: {e}")

def start_bar(textbox_data, from_code, to_code):
    print("In Progress Bar start")
    progressbar = ttk.Progressbar(mode="indeterminate")
    progressbar.grid(row=6, column=15, columnspan=12, sticky='W' + 'E' + 'S')
    progressbar.start()
    #button.destroy()
    print("Starting the thread for translate")
    threading.Thread(target=argo_function, args=(textbox_data, from_code, to_code, progressbar)).start()
    messagebox.showinfo('Processing...', "Translation Inprogress...")


def translate_german():
    # t1 = threading.Thread(target = start_bar)
    textbox_data = (my_text.get(1.0, END))
    print(textbox_data)
    # my_text.delete(1.0, END)
    from_code = "de"
    to_code = "en"
    start_bar(textbox_data, from_code, to_code)


def translate_en_german():
    textbox_data = (my_text.get(1.0, END))
    print(textbox_data)
    # my_text.delete(1.0, END)
    from_code = "en"
    to_code = "de"
    start_bar(textbox_data, from_code, to_code)


def translate_arabic():
    textbox_data = (my_text.get(1.0, END))
    print(textbox_data)
    # my_text.delete(1.0, END)
    from_code = "ar"
    to_code = "en"
    start_bar(textbox_data, from_code, to_code)


def translate_en_arabic():
    textbox_data = (my_text.get(1.0, END))
    print(textbox_data)
    # my_text.delete(1.0, END)
    from_code = "en"
    to_code = "ar"
    start_bar(textbox_data, from_code, to_code)


def translate_turkish():
    textbox_data = (my_text.get(1.0, END))
    print(textbox_data)
    # my_text.delete(1.0, END)
    from_code = "tr"
    to_code = "en"
    start_bar(textbox_data, from_code, to_code)


def translate_romania():
    textbox_data = (my_text.get(1.0, END))
    print(textbox_data)
    # my_text.delete(1.0, END)
    from_code = "ro"
    to_code = "en"
    start_bar(textbox_data, from_code, to_code)


def translate_en_turkish():
    textbox_data = (my_text.get(1.0, END))
    print(textbox_data)
    # my_text.delete(1.0, END)
    from_code = "en"
    to_code = "tr"
    start_bar(textbox_data, from_code, to_code)


# Translate Spanish Function
def translate_spanish():
    textbox_data = (my_text.get(1.0, END))
    print(textbox_data)
    # my_text.delete(1.0, END)
    browse_text.set("Processing...")
    from_code = "es"
    to_code = "en"
    start_bar(textbox_data, from_code, to_code)


def translate_en_spanish():
    textbox_data = (my_text.get(1.0, END))
    print(textbox_data)
    from_code = "en"
    to_code = "es"
    start_bar(textbox_data, from_code, to_code)


# Clear the textbox
def clear_text_box():
    my_text.delete(1.0, END)


# Exit button
def exit_tk():
    root.destroy()


# Save file function
def save_file():
    file_path = filedialog.asksaveasfile(initialdir="C:/gui/",
                                         filetypes=[("Text file", ".txt"), ("HTML file", ".html"), ("PDF file", ".pdf"),
                                                    ("Excel file", ".xlsx"), ("Doc file", ".docx"),
                                                    ("All files", ".*")],
                                         )
    file_text = my_text.get(1.0, END)
    # print(file_text)

    string = file_text
    new_string = string.replace("€", "Euro")

    print(new_string)
    print(file_path)
    file_path.write(new_string)
    # file_path.write(file_text)
    file_path.close()

#open_file =''

# Open our pdf file
def open_pdf():
    global open_file
    browse_text1.set("loading...")
    # Grab the filename pf the pdf file
    open_file = filedialog.askopenfilename(
        initialdir="C:/gui/",
        filetypes=(
            ("PDF files", "*.pdf"),
            ("All Files", "*.*")
        )
    )
    # Check to see if there is a file
    print("File path",open_file)
    if open_file:
        pdf_file = PdfReader(open_file)
        browse_text1.set("Select Pdf")

        page_stuff = ""
        for page in pdf_file.pages:
            page_stuff += page.extract_text() + "\n"

        my_text.delete(1.0, tk.END)

        # Add text to textbox
        my_text.insert(1.0, page_stuff)
        browse_text1.set("Select Pdf")


def open_file():
    global open_file
    browse_text.set("loading...")

    # Grab the filename of the file
    open_file = filedialog.askopenfilename(
        initialdir="C:/gui/",
        filetypes=(
            ("All Files", "*.*"),
            ("Text files", "*.txt"),
            ("PDF files", "*.pdf"),
            ("Word files", "*.docx")
        )
    )

    # Check if a file was selected
    print("File path:", open_file)
    if not open_file:
        browse_text.set("Select Docx")
        return

    page_stuff = ""
    if open_file.endswith('.docx'):
        # Read the DOCX file
        word_file = docx.Document(open_file)
        for paragraph in word_file.paragraphs:
            page_stuff += paragraph.text + "\n"
    elif open_file.endswith('.pdf'):
        # Read the PDF file
        with open(open_file, 'rb') as pdf_file:
            reader = PyPDF2.PdfFileReader(pdf_file)
            for page_num in range(reader.numPages):
                page = reader.getPage(page_num)
                page_stuff += page.extract_text() + "\n"
    elif open_file.endswith('.txt'):
        # Read the TXT file
        with open(open_file, 'r') as txt_file:
            page_stuff = txt_file.read()

    # Clear existing text in the textbox
    my_text.delete(1.0, tk.END)
    # Add extracted text to the textbox
    my_text.insert(1.0, page_stuff)
    browse_text.set("Select Docx")

def word_to_word_Conversion(typeOfFile,progressbar):
    try:
        input_file = open_file
        print("Path and type",input_file,typeOfFile)
        #print("Path and type",input_file,typeOfFile)
        if typeOfFile=="pdf":
            input_file = getallwordfiles("source")
            for file in input_file:
                print("inputFile", file)
                translate_german_to_english(file, typeOfFile)
            #input_file = "C:\\Intel\\AI ML\\NLP\\output\\source_word.docx"
        #print("File path", input_file)
        else:
            translate_german_to_english(input_file, typeOfFile)
        #Success
        print("Word conversion done")
    except Exception as e:
        #Error
        print(e)
    finally:
        progressbar.stop()
        progressbar.destroy()

def create_pdf():
    file_text = my_text.get(1.0, END)
    # print(file_text)

    file_path = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF files", "*.pdf")])

    string = file_text
    new_string = string.replace("€", "Euro")
    print(new_string)

    if not file_path:
        print("Save operation cancelled")
        return

    # print(file_text)
    try:
        pdf = FPDF()
        pdf.add_page()
        # pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font("Arial", size=12)
        # To reduce line spacing
        pdf.set_auto_page_break(auto=True, margin=15)

        pdf.multi_cell(0, 12, txt=new_string, border=0, align='L', fill=False)

        pdf.output(file_path)
        # pdf.output("Translated.pdf")
        messagebox.showinfo("Success", f"Translated PDF saved")
    except Exception as e:
        print(f"Error occured: {e}")


def create_excel():
    try:
        # Retrieve text from your text widget
        file_text = my_text.get(1.0, END)

        # Replace "C" with "Euro" in the text
        new_string = file_text.replace("€", "Euro")
        print(new_string)

        # Create a workbook and add a worksheet
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Translated Text"

        # Split the text into lines and write to the Excel file
        for row, line in enumerate(new_string.split('\n'), start=1):
            ws.cell(row=row, column=1, value=line)
        file_path = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel files", "*.xlsx")])
        if file_path:
            # Save the workbook
            wb.save(file_path)
            messagebox.showinfo("Success", "Translated Excel file saved")
        else:
            messagebox.showwarning("Cancelled", "Save operation cancelled")
        # Save the workbook
        # wb.save("Translated.xlsx")

        # messagebox.showinfo("Success", "Translated Excel file saved")

    except Exception as e:

        messagebox.showerror("Error", f"An error occurred: {e}")


def create_word():
    try:
        # Retrieve text from your text widget
        file_text = my_text.get(1.0, tk.END)

        # Replace "C" with "Euro" in the text string
        file_text = file_text.replace("€", "Euro")
        print(file_text)

        # Create a new Word document
        doc = docx.Document()

        # Add a title to the document
        # doc.add_heading('Translated Text', 0)

        # Split the text into lines and add to the Word document
        for line in file_text.split('\n'):
            doc.add_paragraph(line)

            # Ask user to choose a save location
        save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Documents", "*.docx")])

        if save_path:
            # Save the Word document
            doc.save(save_path)
            messagebox.showinfo("Success", "Translated Word document saved")
        else:
            messagebox.showwarning("Cancelled", "Save operation cancelled.")

    except Exception as e:
        messagebox.showerror("Error", f"An error occurred: {e}")


def progressbar_fnct(translate, progressbar):
    print("inprogress bar", type(progressbar))

    if translate == 1:
        # progressbar.place(x=60, y=90, width=200)
        progressbar.grid(row=6, column=10, columnspan=12, sticky='W' + 'E' + 'S')
        # Start moving the indeterminate progress bar.
        progressbar.start()


    else:
        print("------------stop bar")
        progressbar.grid(row=6, column=8, columnspan=10, sticky='W' + 'E' + 'S')
        type(progressbar)
        progressbar.stop()
        progressbar.place_forget()
        progressbar.destroy()
        progressbar = ttk.Progressbar(mode="indeterminate")

        print("------------------No Selected Option")


def on_select(selection):
    # selection = event.widget.get()
    if selection == "Select Language":
        selected_option.set(options[0])
        return
    command_map = {
        "German-English": translate_german,
        "Spanish-English": translate_spanish,
        "Arabic-English": translate_arabic,
        "Turkish-English": translate_turkish,
        "Romanian-English": translate_romania,
        "English-German": translate_en_german,
        "English-Arabic": translate_en_arabic,
        "English-Spanish": translate_en_spanish,
        "English-Turkish": translate_en_turkish
    }

    command = command_map.get(selection)
    # print("----",command)
    # global translate;
    if command:
        # translate = 1
        # print("Translate",translate)
        command()


options = [
    "Select Language",
    "German-English",
    "Spanish-English",
    "Arabic-English",
    "Turkish-English",
    "Romanian-English",
    "English-German",
    "English-Arabic",
    "English-Spanish",
    "English-Turkish",

]

selected_option = tk.StringVar()
# print("------------",translate)
selected_option.set(options[0])

dropdown = tk.OptionMenu(root, selected_option, *options, command=on_select)
# dropdown.pack(pady=20)
dropdown.config(bg='#77B5BB', fg="black", height=2, width=12, font=("Arial", 10))
dropdown.grid(row=10, column=9)

def getallPdfFiles(value):
    print("Value",value)
    print(pathlib.Path(get_current_directory_path()))
    pdf_doc_lst = list(pathlib.Path(get_current_directory_path()).glob(value+'*.pdf'))
    print("All PDF file",pdf_doc_lst)
    return  pdf_doc_lst

def deleteAllWordFiles():
    docxFiles = getallwordfiles("Result")
    if len(docxFiles) == 0:
        #pop up message "PDF Merge failed"
        messagebox.showerror('Error...',"PDF Merge failed")
        return
    for file in docxFiles:
        os.remove(file)


def deleteALlOldPdfFiles():
    pdffiles = getallPdfFiles("Translated")
    if len(pdffiles) == 0:
        #pop up message "PDF Merge failed"
        messagebox.showerror('Error...',"PDF Merge failed")
        return
    for file in pdffiles:
        os.remove(file)
    print("Deleted previous Pdf file")
'''
def mergerPdf():
    pdffiles = getallPdfFiles("Translated")
    if len(pdffiles) == 0:
        #pop up message "PDF Merge failed"
        messagebox.showerror('Error...',"PDF Merge failed")
        return
    merger = PdfMerger()
    try:
        for pdf in pdffiles:
            merger.append(open(pdf, 'rb'))
        outputpath = "Result_"+getcurrentDateTime()+".pdf"
        with open(outputpath, "wb") as fout:
            merger.write(fout)
        print("PDF Merged Successfully")
        return outputpath
    except Exception as e:
        print("Unable to Merge PDF"+e)
'''

def convertwordToPdf():
    try:
        doc = docDocument()
        translated_docx = getallwordfiles("Result")
        print("Translated docx",translated_docx)
        for docx in translated_docx:
            print("Document",docx)
            doc.LoadFromFile(str(docx))
            doc.Watermark = None
            doc.SaveToFile("Translated_output_"+getcurrentDateTime()+".pdf",docFormat.PDF)
        doc.Close()
    except Exception as e:
        print("Exception as e",e)


def convertPdfToWord(typeOfFile,progressbar):
    try:
        convertSourcePdfToWord()
        word_to_word_Conversion(typeOfFile, progressbar)
        convertwordToPdf()
        #House keeping Activity
        #deleteAllWordFiles()
    except Exception as e:
        #Add error pdf translation
        messagebox.showinfo('Error...', "Exception in PDF Translation!")
        print("Exception in PDF Translation",e)


def startProgressBar(typeOfFile):
    #print("In Progress Bar start")
    progressbar = ttk.Progressbar(mode="indeterminate")
    progressbar.grid(row=6, column=10, columnspan=12, sticky='W' + 'E' + 'S')
    progressbar.start()
    try:
        #print("Starting the thread for translate")

        threading.Thread(target=convertPdfToWord, args=(typeOfFile,progressbar)).start()
        #messagebox.showinfo('Processing...', "Translation is Inprogress...Please click ok to close this window!!!")
    except Exception as e:
        print(e.with_traceback())

def wordstartProgressBar(typeOfFile):
    #print("In Progress Bar start")
    progressbar = ttk.Progressbar(mode="indeterminate")
    progressbar.grid(row=6, column=10, columnspan=12, sticky='W' + 'E' + 'S')
    progressbar.start()
    try:
        if typeOfFile=='word':
            threading.Thread(target=word_to_word_Conversion, args=(typeOfFile,progressbar)).start()

    except Exception as e:
        print(e.with_traceback())


def getallhtmlfile(value):
    #print("Value",value)
    html_doc_lst = list(pathlib.Path(get_current_directory_path() + "\\output\\").glob(value+'*.docx'))
    print("All html file",type(html_doc_lst))
    return  html_doc_lst

def getallwordfiles(value):
    #print("Value",value)
    html_doc_lst = list(pathlib.Path(get_current_directory_path() + "\\output\\").glob(value+'*.docx'))
    print("All Word file",type(html_doc_lst))
    return  html_doc_lst

def convertSourcePdfToWord():
    # Load a pdf document
    #inputfile = doc.LoadFromFile(open_file)
    try:
        doc = PdfDocument()
        messagebox.showinfo('Processing...', "Translation Inprogress! Find the updates in App Directory !")
        doc.LoadFromFile(open_file)
        # Convert to doc file.
        print("open_file",open_file)
        doc.SaveToFile("output/source_word.docx", PdfFIleFormat.DOCX)
        doc.Close()
        print("PDF to word converted successfullyy")
    except Exception as e:
        print("Exception as ",e)
'''
def convertAllHtmlFiles(typeOfFile,progressbar):
    convertToHtml()
    pdfpath =''
    htmlfiles = getallhtmlfile('s')
    try:
        if len(htmlfiles) == 0:
            #messagebox.showerror('Error...', "Docx File Not selected")
            raise FileNotFoundError("No files found for translation")

        for htmlfile in htmlfiles:
            TranslateHtml(htmlfile)
        convertToPDF()

    except:
        pass
    
        
        pdfpath = mergerPdf()
    except FileNotFoundError as F:
        print(F.with_traceback())
    except Exception as e:
        print(e.with_traceback())

    finally:
        progressbar.stop()
        progressbar.destroy()
        #House Keeping activity
        deleteAllOldHtmlFIles()
        deleteALlOldPdfFiles()

    '''
'''
def convertToWord():
    
    htmlfiles = getallhtmlfile('o')
    try:
        for htmlfile in htmlfiles:
            pass
            converter.convert(f'file:///{htmlfile}', 'TranslatedInvoice_' + getcurrentDateTime() + '.docx')
    except Exception as e:
        # Pop up to be add "Translation Failed
        messagebox.showerror('Error...',"Translation Failed")
        print("Error in Word Conversion",e.with_traceback())
        return
    else:
        #Pop up to be add "Translation is done successfully
        messagebox.showinfo('Completed!',"Translation successfully completed")
        print("Translation successfully completed")
'''

def conversionToWord():
    startProgressBar('word')
     #Pop up to be added "Word conversion is successful
    messagebox.showinfo('Completed',"Word conversion is successful")


'''
def convertdocxAllHtmlFiles(typeOfFile,progressbar):
    convertdocxtoHtml()
    pdfpath =''
    htmlfiles = getallhtmlfile('s')
    try:
        if len(htmlfiles) == 0:
            raise FileNotFoundError("No files found for translation")
        for htmlfile in htmlfiles:
            TranslateHtml(htmlfile)
        #convertToword()
        #convertToWord()
        #pdfpath = mergerPdf()
        if typeOfFile == 'word':
            pass
    except FileNotFoundError as F:
        print(F.with_traceback())
    except Exception as e:
        print(e.with_traceback())

    finally:
        progressbar.stop()
        progressbar.destroy()
        #House Keeping activity
        deleteAllOldHtmlFIles()
        deleteALlOldPdfFiles()
'''
def getcurrentDateTime():
    return datetime.datetime.now().strftime("%Y%m%d-%H%M%S")

def deleteAllOldHtmlFIles():
    outputfile = getallhtmlfile('')
    if len(outputfile) == 0:
        return
    for file in outputfile:
        os.remove(file)
    print("Deleted previous HTML file")



def TranslateHtml(htmlfile):
    dict_valuesReplace = {
        "Single price": "Unit price",
        "KUNDEN NR": "Customer No.",
        "USt.":"Vat",
        "RECOVERY": "YOUR BILL",
        "Please worry for sufficient account coverage.": "Please ensure that your account has sufficient funds.",
        "Woman":"Mrs."
    }
    from_code = "de"
    to_code = "en"
    print("Html doument",htmlfile)
    html_content = ""
    with open(htmlfile, 'r', encoding='utf-8') as file:
        html_content += file.read()

    parser = etree.HTMLParser()
    tree = etree.fromstring(html_content, parser)
    elements = tree.xpath("//*[contains(text(),'www*html')]")
    removed_element_info = []
    try:
        for element in elements:
            parent = element.getparent()
            if parent is not None:
                index = parent.index(element)
                removed_element_info.append((element, parent, index))
                parent.remove(element)
            else:
                pass
    except Exception as e:
        print("Eroor in get parent ", e)

    html_content = etree.tostring(tree, pretty_print=True, encoding='unicode')
    response = BeautifulSoup(html_content, 'html.parser')
    # Download and install Argos Translate package
    available_packages = argostranslate.package.get_available_packages()
    available_package = list(
        filter(
            lambda x: x.from_code == from_code and x.to_code == to_code, available_packages
        )
    )[0]
    download_path = available_package.download()
    argostranslate.package.install_from_path(download_path)
    # Translate
    installed_languages = argostranslate.translate.get_installed_languages()
    from_lang = list(filter(lambda x: x.code == from_code, installed_languages))[0]
    to_lang = list(filter(lambda x: x.code == to_code, installed_languages))[0]
    translation = from_lang.get_translation(to_lang)
    translated_soup = translatehtml.translate_html(translation, str(response))
    translated_soup = str(translated_soup)
    for key,value in dict_valuesReplace.items():
        translated_soup = translated_soup.replace(key,value)
    #translated_soup = replace_word(translated_soup,dict_valuesReplace)
    translated_tree = etree.fromstring(translated_soup, parser)
    for element, parent, index in removed_element_info:
        translated_parent = translated_tree.xpath(parent.getroottree().getpath(parent))[0]
        translated_parent.insert(index, element)

    translated_soup = etree.tostring(translated_tree, pretty_print=True, encoding='unicode')

    try:

        output_file = get_current_directory_path() + "\\output\\output_"+getcurrentDateTime()+".html"
        with open(output_file, "w",encoding="utf-8") as file:
            file.write(str(translated_soup))
    except Exception as e:
        #Pop mu message to be added
        messagebox.showerror('Error...',"Unable to create the PDF")
        print("Exception",e)
        print("Unable to create a PDF")
        return
    else:
        print("Translated HTML successfully")

# Load a PDF document
def convertToHtml():
    try:
        messagebox.showinfo('Processing...', "Translation Inprogress! Find the updates in App Directory !")
        doc = PdfDocument()
        print("Convert to html", open_file)
        doc.LoadFromFile(open_file)
        # Set the conversion options to
        convertOptions = doc.ConvertOptions
        # Specify convert options
        convertOptions.SetPdfToHtmlOptions(True, True, 2, True)  # Save the PDF document to HTML format
        doc.SaveToFile("output/source.html", PdfFIleFormat.HTML)
        # Dispose resources
        doc.Dispose()
    except Exception as e:
        print("Error",e)


'''
def convertdocxtoHtml():
    # Load the DOCX file
    doc = docx.Document(open_file)

    # Initialize the conversion options
    convert_options = doc.ConvertOptions

    # Specify the conversion settings (you can adjust these as needed)
    convert_options.SetPdfToHtmlOptions(False, True, 1, False)

    # Save the DOCX document to HTML format
    #doc.SaveToFile("output/source.html", FileFormat.HTML)

    # Dispose of resources
    doc.Dispose()
'''


def conversionToPdf():
    startProgressBar('pdf')


'''
def pdfToWord(pdfpath):
    try:
        word = PdfDocument()
        word.LoadFromFile(pdfpath)
        #word.SaveToFile(get_current_directory_path() + "\\wordDocument\\Translated_"+getcurrentDateTime()+".doc", FileFormat.DOC)
        print("Converted to word successfully")
    except:
        print("Error in Word conversion")
'''

def convertToPDF():
    #Save the PDF
    #htmlfiles = getallhtmlfile('o')
    #html_doc_lst = list(pathlib.Path(get_current_directory_path() + "\\output\\").glob("output*.html"))
    #print(type(html_doc_lst))
    #htmlfiles = [["C:/Users/ShaikhJ11/Downloads/NLP (1)/NLP/output/output_20240911-181233.html"]]

    outputpath = str(pathlib.Path(__file__).parent.resolve()) +"\\output\\"
    print("Output path",outputpath)
    htmlfiles = list(pathlib.Path(outputpath).glob("output*.html"))
    print("Html output FIles",htmlfiles)
    try:
        if(len(htmlfiles)==0):
            messagebox.showerror('Warning...', "No Output File Generated")
            print("No Output files to be processed")
            return
        for htmlfile in htmlfiles:
            #doc = docDocument()
            print("In PDF FUnction",str(htmlfile))
            #doc.LoadFromFile(str(htmlfile), FileFormat.Html, XHTMLValidationType.none)
            #doc.SaveToFile("TranslatedInvoice_" + getcurrentDateTime()+".pdf",FileFormat.PDF)
            converter.convert(f'file:///{htmlfile}', 'TranslatedInvoice_' + getcurrentDateTime() + '.pdf')
            #doc.Close()
    except Exception as e:
        # Pop up to be add "Translation Failed
        messagebox.showerror('Error...',"Translation Failed")
        print("Error in PDF Conversion",e.with_traceback())
        return
    else:
        #Pop up to be add "Translation is done successfully
        messagebox.showinfo('Completed!',"Translation successfully completed")
        print("Translation successfully completed")


def conversionToWord():
     wordstartProgressBar('word')
     #Pop up to be added "Word conversion is successful
     messagebox.showinfo('Processing...', "Translation Inprogress! Find the updates in App Directory !")




def on_select(selection):
    # selection = event.widget.get()
    if selection == "Save as":
        return
    command_map = {
        "Save as": save_file,
        "Pdf": create_pdf,
        "Word": create_word,
        "Text": save_file,
        "Excel": create_excel
    }
    command = command_map.get(selection)
    if command:
        command()


options = [
    "Save as",
    "Pdf",
    "Word",
    "Text",
    "Excel"
]

selected_option = tk.StringVar()
selected_option.set(options[0])

dropdown = tk.OptionMenu(root, selected_option, *options, command=on_select)
# dropdown.pack(pady=20)
dropdown.config(bg='#77B5BB', fg="black", height=2, width=10, font=("Arial", 10))
dropdown.grid(row=10, column=15)


# Function to save text in various formats

def save_text():
    file_type = file_type_var.get()
    if file_type == "pdf":
        create_pdf()
    elif file_type == "xlsx":
        create_excel()
    elif file_type == "txt":
        save_file()
    elif file_type == "docx":
        create_word()


file_type_var = StringVar(root)
file_type_var.set("pdf")  # default value

# Enter button
browse_text1 = StringVar()
button = Button(root, textvariable=browse_text1, command=open_pdf, bg='#AA5D55', fg="black", height=2,
                width=10, font=("Arial", 10))
browse_text1.set("Select Pdf")
# button.grid(row=10, rowspan=2, column=0)
button.grid(row=10, column=1, columnspan=1, padx=10, pady=10)
# lambda:open_pdf()

# Enter button
browse_text = StringVar()
button = Button(root, textvariable=browse_text, command=open_file, bg='#AA5D55', fg="black", height=2,
                width=12, font=("Arial", 10))
browse_text.set("Select Docx")
# button.grid(row=10, rowspan=2, column=0)
button.grid(row=10, column=1, columnspan=10, padx=10, pady=10)
# lambda:open_pdf()

# Clear Button
button_text = StringVar()
button2 = Button(root, text="Clear", bg='#77B5BB', fg="black", height=2, width=8, font=("Arial", 10),
                 command=clear_text_box)
button_text.set("Clear")
button2.grid(row=10, column=18, padx=15, pady=15)

# Exit button
button3 = Button(root, text="Exit", command=exit_tk, bg='#77B5BB', fg="black", height=2, width=8, font=("Arial", 10))
button3.grid(row=10, column=19, padx=15, pady=15)

button4 = Button(root, text="Pdf Invoice", command=conversionToPdf,bg='#77B5BB', fg="black", height=2, width=10, font=("Arial", 10))
button4.grid(row=10, column=17, padx=15, pady=15)


button5 = Button(root, text="Word Invoice", command=conversionToWord, bg='#77B5BB', fg="black", height=2, width=10, font=("Arial", 10))
button5.grid(row=10, column=16, padx=15, pady=15)


root.mainloop()
# configure_network_threaded()
